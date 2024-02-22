import os
import re

import docx


class Templating:
    def __init__(
        self,
        input_file: str = None,
        output_dir: str = None,
        template_start: str = "\[",
        template_end: str = "]",
        templates: dict = {},
    ) -> None:
        """Class for templating .docx files, creates a .docx document based on the input_file
        substituting in the data based on the templates.

        Args:
            input_file (str, optional): Path to the input file. Defaults to None.
            output_dir (str, optional): Path to the output folder. Defaults to None.
            template_start (str, optional): Regex pattern for the template start designator. Defaults to "\[".
            template_end (str, optional): Regex pattern for the template end designator. Defaults to "]".
            templates (dict, optional): Dict for the templates to substitute. Defaults to {}.
        """
        self.document = docx.Document(input_file)
        self.output_dir = output_dir
        self.template_start = template_start
        self.template_end = template_end
        self.templates = templates.copy()
        self.input_elements = self.get_input_elements()

    def get_input_elements(self):
        """Gets the sections and content elements of the .docx file.

        Elements are either docx.paragraph or docx.table"""
        section_components = []
        section_elements = []
        for section in self.document.sections:
            section_components += [
                section.even_page_footer,
                section.even_page_header,
                section.first_page_footer,
                section.first_page_header,
                section.footer,
                section.header,
            ]
            for component in section_components:
                section_elements += component.paragraphs
                section_elements += component.tables

            section_elements += section.iter_inner_content()
        return section_elements

    def sub(self, template: str, value: str):
        """Substitutes TEMPLATE with VALUE in the document.

        Iterates over the input elements checking if the element contains the TEMPLATE.
        If so, it iterates over the runs in the elements and substitutes in the VALUE."""
        reg = re.compile(self.template_start + template + self.template_end)

        for element in self.input_elements:
            try:
                if reg.search(element.text):
                    for i, run in enumerate(element.runs):
                        if reg.search(run.text):
                            run.text = reg.sub(value, run.text)
                        elif re.search(self.template_start + template + "$", run.text):
                            if re.search("^" + self.template_end, element.runs[i + 1].text):
                                run.text = re.sub(
                                    self.template_start + template + "$", value, run.text
                                )
                                element.runs[i + 1].text = re.sub(
                                    "^" + self.template_end, "", element.runs[i + 1].text
                                )
                        elif re.search(self.template_start + "$", run.text):
                            if re.search(
                                "^" + template + self.template_end, element.runs[i + 1].text
                            ):
                                run.text = re.sub(self.template_start + "$", "", run.text)
                                element.runs[i + 1].text = re.sub(
                                    "^" + template + self.template_end,
                                    value,
                                    element.runs[i + 1].text,
                                )
                            elif element.runs[i + 1].text == template and re.search(
                                "^" + self.template_end, element.runs[i + 2].text
                            ):
                                run.text = re.sub(self.template_start + "$", "", run.text)
                                element.runs[i + 1].text = re.sub(
                                    template, value, element.runs[i + 1].text
                                )
                                element.runs[i + 2].text = re.sub(
                                    "^" + self.template_end, "", element.runs[i + 2].text
                                )
            # If the exception id raised the given element does not have a text attribute.
            # That happes if the element is docx.table instead of docx.paragraph.
            except AttributeError:
                for row in element.rows:
                    for cell in row.cells:
                        for cell_p in cell.paragraphs:
                            for i, cell_run in enumerate(cell_p.runs):
                                if reg.search(cell_run.text):
                                    cell_run.text = reg.sub(value, cell_run.text)
                                elif re.search(self.template_start + template + "$", cell_run.text):
                                    if re.search("^" + self.template_end, element.runs[i + 1].text):
                                        cell_run.text = re.sub(
                                            self.template_start + template + "$",
                                            value,
                                            cell_run.text,
                                        )
                                        element.runs[i + 1].text = re.sub(
                                            "^" + self.template_end, "", element.runs[i + 1].text
                                        )
                                elif re.search(self.template_start + "$", cell_run.text):
                                    if re.search(
                                        "^" + template + self.template_end, element.runs[i + 1].text
                                    ):
                                        cell_run.text = re.sub(
                                            self.template_start + "$", "", cell_run.text
                                        )
                                        element.runs[i + 1].text = re.sub(
                                            "^" + template + self.template_end,
                                            value,
                                            element.runs[i + 1].text,
                                        )
                                    elif element.runs[i + 1].text == template and re.search(
                                        "^" + self.template_end, element.runs[i + 2].text
                                    ):
                                        cell_run.text = re.sub(
                                            self.template_start + "$", "", cell_run.text
                                        )
                                        element.runs[i + 1].text = re.sub(
                                            template, value, element.runs[i + 1].text
                                        )
                                        element.runs[i + 2].text = re.sub(
                                            "^" + self.template_end, "", element.runs[i + 2].text
                                        )

    def sub_templates(self):
        """Subtitutes all the templates in the documents."""
        for template, value in self.templates.items():
            self.sub(template=template, value=value)

    def save(self):
        """Save the document as 'output.docx' in the designated output folder.

        Returns the path to the file."""
        save_path = os.path.join("", self.output_dir, "output.docx")
        self.document.save(save_path)
        return save_path
